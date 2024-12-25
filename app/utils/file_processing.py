import pdfplumber
import docx
import io
import streamlit as st
from PyPDF2 import PdfReader

def extract_text_from_file(uploaded_file):
    """Extract text from various file types"""
    try:
        if uploaded_file.type == "application/pdf":
            try:
                # Get PDF bytes
                pdf_bytes = uploaded_file.getvalue()
                pdf_stream = io.BytesIO(pdf_bytes)
                
                # Try pdfplumber first
                try:
                    with pdfplumber.open(pdf_stream) as pdf:
                        text = ""
                        for page in pdf.pages:
                            extracted = page.extract_text()
                            if extracted:
                                text += extracted + "\n"
                        
                        if text.strip():  # If we got valid text
                            return text
                except:
                    # Fallback to PyPDF2 if pdfplumber fails
                    pdf_stream.seek(0)  # Reset stream position
                    reader = PdfReader(pdf_stream)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    
                    if not text.strip():  # If still no valid text
                        raise Exception("Could not extract text from PDF")
                    
                    return text
                
                # Debug info
                with st.expander("🔍 Debug Information", expanded=False):
                    st.markdown(f"""
                        <p class='debug-text'>PDF Extraction:</p>
                        <p class='debug-text'>- Characters extracted: {len(text)}</p>
                        <p class='debug-text'>- Preview: {text[:200]}</p>
                    """, unsafe_allow_html=True)
                
            except Exception as e:
                st.error(f"PDF extraction error: {str(e)}")
                return None
                
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                # Read DOCX
                doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                
                # Debug info
                with st.expander("🔍 Debug Information", expanded=False):
                    st.markdown(f"""
                        <p class='debug-text'>DOCX Extraction:</p>
                        <p class='debug-text'>- Paragraphs: {len(doc.paragraphs)}</p>
                        <p class='debug-text'>- Characters: {len(text)}</p>
                        <p class='debug-text'>- Preview: {text[:200]}</p>
                    """, unsafe_allow_html=True)
                
                return text
                
            except Exception as e:
                st.error(f"DOCX extraction error: {str(e)}")
                return None
                
        elif uploaded_file.type == "text/plain":
            try:
                # Read TXT
                text = uploaded_file.getvalue().decode('utf-8')
                
                # Debug info
                with st.expander("🔍 Debug Information", expanded=False):
                    st.markdown(f"""
                        <p class='debug-text'>TXT Extraction:</p>
                        <p class='debug-text'>- Characters: {len(text)}</p>
                        <p class='debug-text'>- Preview: {text[:200]}</p>
                    """, unsafe_allow_html=True)
                
                return text
                
            except Exception as e:
                st.error(f"TXT extraction error: {str(e)}")
                return None
                
        else:
            st.error(f"Unsupported file type: {uploaded_file.type}")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return None

def extract_text_from_docx(docx_file) -> Optional[str]:
    """
    Extract text content from a DOCX file.
    
    Args:
        docx_file: Uploaded DOCX file from Streamlit
        
    Returns:
        str: Extracted text content or None if extraction fails
    """
    try:
        text_content = docx2txt.process(docx_file)
        return text_content
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return None

def process_resume_file(uploaded_file) -> Optional[Tuple[str, str, str]]:
    """
    Process an uploaded resume file and extract its content.
    
    Args:
        uploaded_file: File uploaded through Streamlit's file_uploader
        
    Returns:
        tuple: (filename, content, file_type) or None if processing fails
    """
    try:
        if not uploaded_file:
            st.error("No file was uploaded")
            return None
            
        # Get filename without extension
        file_name = uploaded_file.name.rsplit('.', 1)[0]
        file_type = uploaded_file.type
        
        # Process different file types
        if file_type == "application/pdf":
            content = extract_text_from_file(uploaded_file)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content = extract_text_from_docx(uploaded_file)
        elif file_type in ["text/plain", "application/text"]:
            content = uploaded_file.getvalue().decode('utf-8', errors='replace')
        else:
            st.error(f"Unsupported file type: {file_type}")
            return None

        if content:
            return file_name, content, file_type
        else:
            st.error(f"Could not extract content from {uploaded_file.name}")
            return None
            
    except Exception as e:
        st.error(f"Error processing file {uploaded_file.name}: {str(e)}")
        return None

def validate_file_size(file, max_size_mb: int = 10) -> bool:
    """
    Validate if the uploaded file is within size limits.
    
    Args:
        file: Uploaded file
        max_size_mb: Maximum allowed file size in MB
        
    Returns:
        bool: True if file size is valid, False otherwise
    """
    try:
        file_size_mb = len(file.getvalue()) / (1024 * 1024)
        if file_size_mb > max_size_mb:
            st.error(f"File size ({file_size_mb:.1f}MB) exceeds the {max_size_mb}MB limit")
            return False
        return True
    except Exception as e:
        st.error(f"Error checking file size: {str(e)}")
        return False

def validate_file_type(file, allowed_types: tuple = ('pdf', 'docx', 'txt')) -> bool:
    """
    Validate if the uploaded file is of an allowed type.
    
    Args:
        file: Uploaded file
        allowed_types: Tuple of allowed file extensions
        
    Returns:
        bool: True if file type is valid, False otherwise
    """
    try:
        if '.' not in file.name:
            st.error("File has no extension")
            return False
            
        file_ext = file.name.rsplit('.', 1)[1].lower()
        if file_ext not in allowed_types:
            st.error(f"File type '.{file_ext}' is not supported. Please upload {', '.join(allowed_types)}")
            return False
            
        # Additional MIME type validation
        mime_types = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': ['text/plain', 'application/text']
        }
        
        if file_ext in mime_types:
            expected_mime = mime_types[file_ext]
            if isinstance(expected_mime, list):
                if file.type not in expected_mime:
                    st.error(f"Invalid file content for .{file_ext} file")
                    return False
            elif file.type != expected_mime:
                st.error(f"Invalid file content for .{file_ext} file")
                return False
                
        return True
    except Exception as e:
        st.error(f"Error checking file type: {str(e)}")
        return False